import os
import sys
import json
from typing import List, Any

import pandas as pd
from docx import Document as DocxDocument
from pypdf import PdfReader
from langchain_core.documents import Document

from src.exception import RagException
from src.logger import logger


def clean_value(value: Any) -> str:
    """
    Convert a dataframe value into a clean string.
    """
    if pd.isna(value):
        return ""

    return str(value).strip()


def clean_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """
    Standardize dataframe column names and replace missing values.
    """
    df = df.copy()
    df.columns = [str(col).strip() for col in df.columns]
    df = df.fillna("")

    return df


def create_table_schema_document(
    file_path: str,
    file_name: str,
    file_type: str,
    df: pd.DataFrame,
    sheet_name: str | None = None,
) -> Document:
    """
    Create a schema-level document for a table.

    This helps the retriever understand available columns,
    row count, and table structure.
    """
    schema_text = f"""
File Name: {file_name}
File Type: {file_type}
Sheet Name: {sheet_name if sheet_name else "N/A"}
Total Rows: {len(df)}
Total Columns: {len(df.columns)}

Columns:
{", ".join([str(col) for col in df.columns])}

Column Details:
""" + "\n".join([f"- {col}" for col in df.columns])

    metadata = {
        "source": file_path,
        "file_name": file_name,
        "file_type": f"{file_type}_schema",
    }

    if sheet_name:
        metadata["sheet_name"] = sheet_name

    return Document(
        page_content=schema_text.strip(),
        metadata=metadata,
    )


def create_row_documents(
    file_path: str,
    file_name: str,
    file_type: str,
    df: pd.DataFrame,
    sheet_name: str | None = None,
) -> List[Document]:
    """
    Convert each dataframe row into a separate document.

    Row-level documents improve retrieval for entity,
    record, and ID-based questions.
    """
    documents = []

    for index, row in df.iterrows():
        row_details = "\n".join(
            [f"{col}: {clean_value(row[col])}" for col in df.columns]
        )

        row_text = f"""
File Name: {file_name}
File Type: {file_type}
Sheet Name: {sheet_name if sheet_name else "N/A"}
Row Number: {index + 1}

{row_details}
""".strip()

        metadata = {
            "source": file_path,
            "file_name": file_name,
            "file_type": file_type,
            "row": index + 1,
        }

        if sheet_name:
            metadata["sheet_name"] = sheet_name

        documents.append(
            Document(
                page_content=row_text,
                metadata=metadata,
            )
        )

    return documents


def create_json_documents(
    file_path: str,
    file_name: str,
    data: Any,
) -> List[Document]:
    """
    Convert JSON content into retrievable documents.

    FAQ-style JSON files are indexed as separate
    question-answer documents.
    """
    documents = []

    if isinstance(data, dict) and isinstance(data.get("faqs"), list):
        company = data.get("company", "")
        document_type = data.get("document_type", "")
        version = data.get("version", "")
        last_updated = data.get("last_updated", "")

        for faq in data["faqs"]:
            question = faq.get("question", "")
            answer = faq.get("answer", "")
            category = faq.get("category", "")
            faq_id = faq.get("id", "")

            faq_text = f"""
File Name: {file_name}
Company: {company}
Document Type: {document_type}
Version: {version}
Last Updated: {last_updated}
Category: {category}
Question: {question}
Answer: {answer}
""".strip()

            if faq_text:
                documents.append(
                    Document(
                        page_content=faq_text,
                        metadata={
                            "source": file_path,
                            "file_name": file_name,
                            "file_type": "json",
                            "category": category,
                            "faq_id": faq_id,
                        },
                    )
                )

        return documents

    text = json.dumps(data, indent=2, ensure_ascii=False)

    if text.strip():
        documents.append(
            Document(
                page_content=(
                    f"File Name: {file_name}\n"
                    f"File Type: json\n\n"
                    f"{text.strip()}"
                ),
                metadata={
                    "source": file_path,
                    "file_name": file_name,
                    "file_type": "json",
                },
            )
        )

    return documents


def load_pdf(file_path: str, file_name: str) -> List[Document]:
    """
    Load each PDF page as a separate document.
    """
    documents = []
    reader = PdfReader(file_path)

    for page_no, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""

        if text.strip():
            documents.append(
                Document(
                    page_content=f"""
File Name: {file_name}
File Type: pdf
Page Number: {page_no}

{text.strip()}
""".strip(),
                    metadata={
                        "source": file_path,
                        "file_name": file_name,
                        "file_type": "pdf",
                        "page": page_no,
                    },
                )
            )

    return documents


def load_csv(file_path: str, file_name: str) -> List[Document]:
    """
    Load a CSV file as schema and row-level documents.
    """
    documents = []

    df = pd.read_csv(file_path)
    df = clean_dataframe(df)

    documents.append(
        create_table_schema_document(
            file_path=file_path,
            file_name=file_name,
            file_type="csv",
            df=df,
        )
    )

    documents.extend(
        create_row_documents(
            file_path=file_path,
            file_name=file_name,
            file_type="csv",
            df=df,
        )
    )

    return documents


def load_excel(file_path: str, file_name: str) -> List[Document]:
    """
    Load every Excel sheet as schema and row-level documents.
    """
    documents = []

    excel_data = pd.read_excel(file_path, sheet_name=None)

    for sheet_name, df in excel_data.items():
        df = clean_dataframe(df)

        documents.append(
            create_table_schema_document(
                file_path=file_path,
                file_name=file_name,
                file_type="excel",
                df=df,
                sheet_name=sheet_name,
            )
        )

        documents.extend(
            create_row_documents(
                file_path=file_path,
                file_name=file_name,
                file_type="excel",
                df=df,
                sheet_name=sheet_name,
            )
        )

    return documents


def load_docx(file_path: str, file_name: str) -> List[Document]:
    """
    Load DOCX paragraphs and tables as separate documents.
    """
    documents = []
    doc = DocxDocument(file_path)

    for para_no, para in enumerate(doc.paragraphs, start=1):
        text = para.text.strip()

        if text:
            documents.append(
                Document(
                    page_content=f"""
File Name: {file_name}
File Type: docx
Paragraph Number: {para_no}

{text}
""".strip(),
                    metadata={
                        "source": file_path,
                        "file_name": file_name,
                        "file_type": "docx",
                        "paragraph": para_no,
                    },
                )
            )

    for table_no, table in enumerate(doc.tables, start=1):
        rows = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]

            if any(cells):
                rows.append(" | ".join(cells))

        if rows:
            table_text = "\n".join(rows)

            documents.append(
                Document(
                    page_content=f"""
File Name: {file_name}
File Type: docx
Table Number: {table_no}

{table_text}
""".strip(),
                    metadata={
                        "source": file_path,
                        "file_name": file_name,
                        "file_type": "docx_table",
                        "table": table_no,
                    },
                )
            )

    return documents


def load_txt(file_path: str, file_name: str) -> List[Document]:
    """
    Load a TXT file as a single document.
    """
    documents = []

    with open(file_path, "r", encoding="utf-8") as file:
        text = file.read().strip()

    if text:
        documents.append(
            Document(
                page_content=f"""
File Name: {file_name}
File Type: txt

{text}
""".strip(),
                metadata={
                    "source": file_path,
                    "file_name": file_name,
                    "file_type": "txt",
                },
            )
        )

    return documents


def load_json(file_path: str, file_name: str) -> List[Document]:
    """
    Load a JSON file and convert it into documents.
    """
    with open(file_path, "r", encoding="utf-8") as file:
        data = json.load(file)

    return create_json_documents(
        file_path=file_path,
        file_name=file_name,
        data=data,
    )


def load_documents(folder_path: str = "data/documents") -> List[Document]:
    """
    Load all supported files from a folder as LangChain Document objects.
    """
    documents = []

    try:
        logger.info(f"Started loading documents from: {folder_path}")

        if not os.path.exists(folder_path):
            raise FileNotFoundError(f"Document folder not found: {folder_path}")

        for root, _, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                file_ext = os.path.splitext(file)[1].lower()

                logger.info(f"Loading file: {file_path}")

                try:
                    if file_ext == ".pdf":
                        file_documents = load_pdf(file_path, file)

                    elif file_ext == ".csv":
                        file_documents = load_csv(file_path, file)

                    elif file_ext in [".xlsx", ".xls"]:
                        file_documents = load_excel(file_path, file)

                    elif file_ext == ".docx":
                        file_documents = load_docx(file_path, file)

                    elif file_ext == ".txt":
                        file_documents = load_txt(file_path, file)

                    elif file_ext == ".json":
                        file_documents = load_json(file_path, file)

                    else:
                        logger.warning(f"Unsupported file skipped: {file_path}")
                        continue

                    documents.extend(file_documents)

                    logger.info(
                        f"Successfully loaded file: {file_path} | "
                        f"Docs created: {len(file_documents)}"
                    )

                except Exception as file_error:
                    logger.error(
                        f"Failed to load file {file_path}: {str(file_error)}"
                    )
                    continue

        logger.info(
            f"All documents loaded successfully. Total docs: {len(documents)}"
        )

        return documents

    except Exception as e:
        logger.error(f"Document loading failed: {str(e)}")
        raise RagException(str(e), sys)


if __name__ == "__main__":
    docs = load_documents()

    print(f"Total Documents: {len(docs)}")

    for i, doc in enumerate(docs[:10], start=1):
        print("\n" + "=" * 80)
        print(f"DOC {i}")
        print("=" * 80)
        print(doc.metadata)
        print(doc.page_content[:1000])